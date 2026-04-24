from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1.2)
    section.bottom_margin = Inches(1.2)
    section.left_margin = Inches(1.3)
    section.right_margin = Inches(1.3)

def heading(text, level=1, center=False):
    p = doc.add_heading(text, level=level)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def body(text, bold=False, italic=False, center=False, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(12)
    return p

def quote(text):
    p = doc.add_paragraph(style='Intense Quote')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    run.font.size = Pt(12)
    run.italic = True
    return p

def divider():
    doc.add_paragraph("— · —").alignment == WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("— · —").bold = False
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

# ── Title ──
doc.add_paragraph()
title = doc.add_heading("Confirmación del Matrimonio", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Omar y Natalia")
r.bold = True
r.font.size = Pt(16)

doc.add_paragraph()
divider()
doc.add_paragraph()

# ── Opening ──
body("Yanira, familia, amigos queridos...", italic=True, center=True, space_after=12)

body("Gracias por reunirnos hoy. Me pidieron decir unas pocas palabras, y lo haré desde el corazón, porque lo que este día representa no cabe en discursos largos. Cabe en una sola pregunta:")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("¿Por qué seguir juntos?")
r.bold = True
r.font.size = Pt(13)
p.paragraph_format.space_after = Pt(10)

body("No en los días de sol. En esos es fácil. El amor se demuestra cuando el sol se oculta.")

# ── Proverb ──
doc.add_paragraph()
quote('"El amor verdadero no es encontrar a alguien perfecto,\nsino ver perfección en alguien imperfecto."')
doc.add_paragraph()

body("Y yo añado: el amor verdadero no vive en los momentos cómodos. Vive en los momentos que los pondrán a prueba. Esos días donde la vida les diga que es más fácil soltar que sostener. Esos días donde estarán cansados, o asustados, o perdidos.")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.add_run("En esos días, el amor no es un sentimiento. Es una decisión.")
p.runs[0].bold = True
p.runs[0].font.size = Pt(12)

divider()

# ── Inception ──
doc.add_paragraph()
body("Omar, hay una película que tú conoces.", italic=True)
body("Inception.")

body("En ella, hay una escena que no se olvida. Un tren. Mal le dice a Cobb:")

doc.add_paragraph()
quote('"Estás esperando un tren. Un tren que te llevará muy lejos.\nSabes adónde esperas que te lleve ese tren, pero no estás seguro.\nPero eso no importa... porque estaremos juntos."')
doc.add_paragraph()

body("Yo he pensado mucho en esa frase.")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("El matrimonio es ese tren.")
r.bold = True
r.font.size = Pt(13)
p.paragraph_format.space_after = Pt(10)

body("No saben exactamente a dónde los llevará. Nadie lo sabe. La vida tiene desvíos, túneles, tramos lentos y tramos que van tan rápido que apenas pueden respirar. Habrá estaciones que no esperaban. Momentos en que la ruta cambia sin avisarles.")

body("Pero eso no importa.")

body("Lo que importa es que van juntos. No es el destino lo que define el matrimonio. Es la compañía. Es la mano que encuentras cuando extiendes la tuya en la oscuridad y sabes, sin ver, que ahí va a estar.")

divider()

# ── Direct words ──
doc.add_paragraph()
body("Natalia, Omar...", italic=True, bold=True, center=True)
doc.add_paragraph()

body("El amor que los trae aquí hoy no es una promesa de que todo será fácil.")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = p.add_run("Es una promesa de que ninguna cosa difícil la enfrentarán solos.")
r.bold = True
r.font.size = Pt(12)
p.paragraph_format.space_after = Pt(8)

body("Que cuando el tren vaya lento, esperarán juntos.")
body("Que cuando el camino no esté claro, lo buscarán juntos.")
body("Que cuando uno quiera bajarse, el otro le tomará la mano y dirá:")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('"Quédate. Yo estoy aquí."')
r.italic = True
r.font.size = Pt(12)
p.paragraph_format.space_after = Pt(10)

body("Que este día no sea el más feliz de sus vidas. Que sea el primero de muchos días felices que construirán juntos.")

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Y cuando lleguen los días duros, acuérdense del tren.")
r.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("No importa a dónde va.")
r.font.size = Pt(12)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Importa que van juntos.")
r.bold = True
r.font.size = Pt(14)

doc.add_paragraph()
body("Que Dios los bendiga hoy, mañana, y en todos los desvíos que vengan.", italic=True, center=True)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Amén.")
r.bold = True
r.font.size = Pt(13)

divider()

# ── Biblical quote ──
doc.add_paragraph()
quote('"Muchas aguas no podrán apagar el amor,\nni los ríos lo ahogarán."\n— Cantares 8:7')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Guion preparado con amor, por petición de Yanira.")
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.save("/home/user/ruflo/docs/boda-omar-natalia.docx")
print("Documento creado: docs/boda-omar-natalia.docx")
